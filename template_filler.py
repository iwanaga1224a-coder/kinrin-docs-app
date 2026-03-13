# -*- coding: utf-8 -*-
"""公式様式テンプレート流し込みエンジン

各区の公式様式（Word/Excel）にデータを流し込んで出力する。
対応していない区は doc_generator.py のフォールバック生成を使用。

方式:
  Word (.docx): テーブル内のラベルテキストを検索 → 隣接する空セルに値を書き込む
  Excel (.xlsx): セル座標を直接指定して値を書き込む
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

try:
    import openpyxl
except ImportError:
    openpyxl = None


TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")


# ========== ユーティリティ ==========

def _set_cell_text(cell, text):
    """Wordテーブルのセルにテキストを設定（既存書式をなるべく保持）"""
    if not cell.paragraphs:
        return
    p = cell.paragraphs[0]
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = str(text)
    else:
        run = p.add_run(str(text))
        run.font.name = "游ゴシック"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
        run.font.size = Pt(10)


def _replace_in_cell(cell, old_text, new_text):
    """セル内のテキストを部分置換"""
    for p in cell.paragraphs:
        full_text = p.text
        if old_text in full_text:
            combined = "".join(r.text for r in p.runs)
            replaced = combined.replace(old_text, new_text, 1)
            for i, run in enumerate(p.runs):
                if i == 0:
                    run.text = replaced
                else:
                    run.text = ""
            return True
    return False


def _cell_is_empty(cell):
    """セルが空（全角スペース・改行のみ含む）かどうか"""
    t = cell.text.strip().replace("\u3000", "").replace("\n", "")
    return not t


def _find_value_cell_in_row(row, label_text):
    """行内でラベルテキストを探し、その隣の空セルを返す

    マージされたセルでは同じテキストが複数セルに出現するため、
    ラベルと異なるテキスト（空のセル）を探す。
    """
    cells = list(row.cells)
    # まずラベルが含まれるセルの範囲を特定
    label_end = -1
    for i, cell in enumerate(cells):
        if label_text in cell.text:
            label_end = i

    if label_end < 0:
        return None

    # ラベルの後の空セルを探す
    for i in range(label_end + 1, len(cells)):
        cell = cells[i]
        # ラベルと同じテキスト（マージの繰り返し）はスキップ
        if label_text in cell.text:
            continue
        # 「電話」「㎡」などの固定テキストはスキップ
        t = cell.text.strip()
        if t and not t.startswith("電") and t != "㎡" and "m2" not in t:
            continue
        if _cell_is_empty(cell):
            return cell

    return None


def _find_row_by_label(table, label_pattern):
    """テーブル内でラベルパターンにマッチする行を返す"""
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            if re.search(label_pattern, cell.text):
                return ri, row
    return None, None


# ========== Word (.docx) 汎用フィラー ==========

# 各区の標識設置届のフィールド定義
# key: ward_name, value: {
#   "table_index": テーブル番号,
#   "labels": [(ラベル検索パターン, データキー), ...],
#   "replacements": [(セル内の古いテキスト, データキー), ...]
# }

SIGN_NOTICE_CONFIG = {
    "品川": {
        "table_index": 0,
        "labels": [
            ("１建築物の名称|１\u3000建築物の名称", "building_name"),
            ("２設計者|２\u3000設計者", "designer_info"),
            ("３施工者|３\u3000施工者", "constructor_info"),
            ("５\u3000主要用途|５ 主要用途", "building_use"),
            ("６\u3000工事種別|６ 工事種別", "construction_type"),
            ("８\u3000高.*さ|８ 高.*さ", "height"),
            ("11\u3000着工予定|11 着工予定", "start_date"),
            ("12\u3000完了予定|12 完了予定", "end_date"),
            ("13\u3000連絡先|13 連絡先", "contact_info"),
        ],
    },
    "渋谷": {
        "table_index": 0,
        "labels": [
            ("１\u3000建築物の名称", "building_name"),
            ("２\u3000設計者", "designer_info"),
            ("３\u3000施工者", "constructor_info"),
            ("４\u3000建築計画に関する連絡先", "contact_info"),
            ("７\u3000主.*用.*途", "building_use"),
            ("１３\u3000着工予定", "start_date"),
        ],
    },
    "墨田": {
        "table_index": 0,
        "labels": [
            ("1\u3000建築物の名称", "building_name"),
            ("2\u3000設計者", "designer_info"),
            ("3\u3000施工者", "constructor_info"),
            ("5\u3000主要用途", "building_use"),
        ],
    },
    "江東": {
        "table_index": 0,
        "labels": [
            ("１\u3000建築物の名称", "building_name"),
            ("２\u3000設計者", "designer_info"),
            ("３\u3000施工者", "constructor_info"),
            ("５\u3000主.*用.*途", "building_use"),
        ],
    },
    "北": {
        "table_index": 1,
        "labels": [
            ("建築物の名称", "building_name"),
            ("設計者住所", "designer_info"),
            ("施工者住所", "constructor_info"),
            ("主要用途", "building_use"),
        ],
    },
    "足立": {
        "table_index": 1,
        "labels": [
            ("１\u3000建築物の名称", "building_name"),
            ("２\u3000設計者", "designer_info"),
            ("３\u3000施工者", "constructor_info"),
            ("５\u3000主.*用.*途", "building_use"),
        ],
    },
    "世田谷": {
        "table_index": 0,
        "labels": [
            ("建築物の名称", "building_name"),
            ("設計者", "designer_info"),
            ("施工者", "constructor_info"),
            ("主要用途|用途", "building_use"),
            ("工事種別", "construction_type"),
            ("高.*さ", "height"),
            ("構.*造", "structure"),
            ("基.*礎.*工", "foundation"),
            ("着工予定", "start_date"),
            ("完了予定|工事完了", "end_date"),
            ("連絡先", "contact_info"),
        ],
    },
    "杉並": {
        "table_index": 0,
        "labels": [
            ("建築物の名称", "building_name"),
            ("設計者", "designer_info"),
            ("施工者", "constructor_info"),
            ("主要用途", "building_use"),
            ("工事種別", "construction_type"),
            ("高.*さ", "height"),
            ("構.*造", "structure"),
            ("基.*礎.*工", "foundation"),
            ("着工予定", "start_date"),
            ("完了予定", "end_date"),
            ("連絡先", "contact_info"),
        ],
    },
    "板橋": {
        "table_index": 0,
        "labels": [
            ("建築物.*名称", "building_name"),
            ("設計者", "designer_info"),
            ("施工者", "constructor_info"),
            ("主要用途|用途", "building_use"),
            ("工事種別", "construction_type"),
            ("高.*さ", "height"),
            ("構.*造", "structure"),
            ("基.*礎.*工", "foundation"),
            ("着工予定", "start_date"),
            ("完了予定", "end_date"),
            ("連絡先", "contact_info"),
        ],
    },
    "江戸川": {
        "table_index": 0,
        "labels": [
            ("建築物.*名称", "building_name"),
            ("設計者", "designer_info"),
            ("施工者", "constructor_info"),
            ("主要用途|用途", "building_use"),
            ("工事種別", "construction_type"),
            ("高.*さ", "height"),
            ("構.*造", "structure"),
            ("基.*礎.*工", "foundation"),
            ("着工予定", "start_date"),
            ("完了予定", "end_date"),
            ("連絡先", "contact_info"),
        ],
    },
    "港": {
        "table_index": 0,
        "labels": [
            ("建築物の名称", "building_name"),
            ("設計者", "designer_info"),
            ("施工者", "constructor_info"),
            ("着工予定", "start_date"),
            ("完了予定", "end_date"),
            ("連絡先", "contact_info"),
        ],
    },
    "目黒": {
        "table_index": 0,
        "labels": [
            ("建築物の名称", "building_name"),
            ("設計者", "designer_info"),
            ("施工者", "constructor_info"),
            ("主要用途", "building_use"),
            ("工事種別", "construction_type"),
            ("高.*さ", "height"),
            ("構.*造", "structure"),
            ("基.*礎.*工", "foundation"),
            ("着工予定", "start_date"),
            ("完了予定", "end_date"),
            ("標識設置日", "sign_install_date"),
        ],
    },
    "荒川": {
        "table_index": 0,
        "labels": [
            ("建築物の名称", "building_name"),
            ("設計者", "designer_info"),
            ("施工者", "constructor_info"),
            ("主要用途", "building_use"),
            ("工事種別", "construction_type"),
            ("高.*さ", "height"),
            ("構.*造", "structure"),
            ("基.*礎.*工", "foundation"),
            ("連絡先", "contact_info"),
        ],
    },
}

REPORT_CONFIG = {
    "品川": {
        "table_index": 0,
        "labels": [
            ("建築物の名称", "building_name"),
        ],
    },
    "渋谷": {
        "table_index": 0,
        "labels": [
            ("１\u3000建築物の名称", "building_name"),
        ],
    },
    "墨田": {
        "table_index": 0,
        "labels": [
            ("建築物の名称", "building_name"),
        ],
    },
    "江東": {
        "table_index": 0,
        "labels": [
            ("建築物の名称", "building_name"),
        ],
    },
    "世田谷": {
        "table_index": 0,
        "labels": [
            ("建築物の名称", "building_name"),
        ],
    },
    "杉並": {
        "table_index": 0,
        "labels": [
            ("建築物の名称", "building_name"),
        ],
    },
    "江戸川": {
        "table_index": 0,
        "labels": [
            ("建築物.*名称", "building_name"),
        ],
    },
    "目黒": {
        "table_index": 0,
        "labels": [
            ("建築物の名称", "building_name"),
        ],
    },
}


def _prepare_data(data):
    """入力データから流し込み用のデータを準備

    テンプレートが必要とする全フィールドを網羅する。
    複合フィールド（designer_info等）は元データから組み立てる。
    """
    designer_name = data.get("designer_name", "")
    designer_tel = data.get("designer_tel", "")
    constructor_name = data.get("constructor_name", "")
    constructor_tel = data.get("constructor_tel", "")
    site_manager = data.get("site_manager", "")

    # 複合フィールド
    designer_info = designer_name
    if designer_tel:
        designer_info = f"{designer_name}  TEL: {designer_tel}"

    constructor_info = constructor_name
    if constructor_tel:
        constructor_info = f"{constructor_name}  TEL: {constructor_tel}"

    contact_info = site_manager
    if constructor_tel:
        contact_info = f"{site_manager}  TEL: {constructor_tel}"

    # 階数フォーマット
    floors_above = data.get("floors_above", "")
    floors_below = data.get("floors_below", "")
    floors_text = ""
    if floors_above or floors_below:
        floors_text = f"地上{floors_above}階"
        if floors_below:
            floors_text += f" 地下{floors_below}階"

    return {
        # === 建物基本情報 ===
        "building_name": data.get("building_name", ""),
        "site_name": data.get("site_name", ""),
        "site_address": data.get("site_address", ""),
        "land_number": data.get("land_number", ""),       # 地名地番
        "building_use": data.get("building_use", ""),
        "structure": data.get("structure", ""),
        "foundation": data.get("foundation", ""),          # 基礎工法
        "construction_type": data.get("construction_type", ""),  # 工事種別
        "height": data.get("height", ""),
        "floors_above": floors_above,
        "floors_below": floors_below,
        "floors_text": floors_text,
        "unit_count": data.get("unit_count", ""),          # 総戸数
        "oneroom_count": data.get("oneroom_count", ""),    # ワンルーム戸数
        "site_area": data.get("site_area", ""),
        "building_area": data.get("building_area", ""),
        "total_floor_area": data.get("total_floor_area", ""),
        "zoning": data.get("zoning", ""),                  # 用途地域
        "fire_zone": data.get("fire_zone", ""),            # 防火地域
        "other_zone": data.get("other_zone", ""),          # その他地域地区
        # === 工期 ===
        "start_date": data.get("start_date", ""),
        "end_date": data.get("end_date", ""),
        "work_hours": data.get("work_hours", ""),
        "holidays": data.get("holidays", ""),
        "work_content": data.get("work_content", ""),
        # === 届出情報 ===
        "submit_date": data.get("submit_date", ""),
        "sign_install_date": data.get("sign_install_date", ""),
        "sign_location": data.get("sign_location", ""),
        # === 関係者 ===
        "applicant_name": data.get("applicant_name", ""),
        "applicant_address": data.get("applicant_address", ""),
        "applicant_tel": data.get("applicant_tel", ""),
        "client_name": data.get("client_name", ""),
        "designer_name": designer_name,
        "designer_info": designer_info,
        "designer_tel": designer_tel,
        "constructor_name": constructor_name,
        "constructor_info": constructor_info,
        "constructor_tel": constructor_tel,
        "site_manager": site_manager,
        "contact_info": contact_info,
        # === 説明報告 ===
        "explanation_date": data.get("explanation_date", ""),
        "explanation_method": data.get("explanation_method", ""),
        "target_count": data.get("target_count", ""),
        "explained_count": data.get("explained_count", ""),
        "unexplained_count": data.get("unexplained_count", ""),
        "opinions": data.get("opinions", ""),
    }


def _find_template_file(ward_name, keywords):
    """テンプレートファイルを検索"""
    ward_dir = os.path.join(TEMPLATES_DIR, f"{ward_name}区")
    if not os.path.isdir(ward_dir):
        return None

    # まず .docx を優先
    for fname in os.listdir(ward_dir):
        if not fname.endswith(".docx"):
            continue
        for kw in keywords:
            if kw in fname:
                return os.path.join(ward_dir, fname)

    # 一括ファイル
    for fname in os.listdir(ward_dir):
        if fname.endswith(".docx") and ("様式" in fname or "申請書類" in fname):
            return os.path.join(ward_dir, fname)

    return None


def _remove_seal_marks(doc):
    """文書内の「印」（押印マーク）を削除する"""
    # 段落内
    for p in doc.paragraphs:
        full = "".join(r.text for r in p.runs)
        if "印" in full:
            cleaned = full.replace("　印", "").replace("印", "")
            for i, run in enumerate(p.runs):
                run.text = cleaned if i == 0 else ""
    # テーブル内
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full = "".join(r.text for r in p.runs)
                    if "印" in full:
                        cleaned = full.replace("　印", "").replace("印", "")
                        for i, run in enumerate(p.runs):
                            run.text = cleaned if i == 0 else ""


def _fill_docx_by_labels(template_path, config, data, output_path):
    """ラベル検索方式でWordテンプレートにデータを流し込む"""
    doc = Document(template_path)
    fill_data = _prepare_data(data)

    table_idx = config["table_index"]
    if table_idx >= len(doc.tables):
        return None

    table = doc.tables[table_idx]

    for label_pattern, field_name in config["labels"]:
        value = fill_data.get(field_name, "")
        if not value:
            continue

        # テーブル内でラベルを探す
        for ri, row in enumerate(table.rows):
            found = False
            for cell in row.cells:
                if re.search(label_pattern, cell.text):
                    found = True
                    break

            if not found:
                continue

            # この行内で空セルを探して値を書き込む
            cells = list(row.cells)
            # ラベルセルの範囲を特定（マージで同じテキストが繰り返される）
            label_cells = set()
            for ci, cell in enumerate(cells):
                if re.search(label_pattern, cell.text):
                    label_cells.add(ci)

            # ラベルでない空セルに値を書き込む
            wrote = False
            for ci, cell in enumerate(cells):
                if ci in label_cells:
                    continue
                # 前のセルと同じ内容ならマージされたセル、スキップ
                if ci > 0 and cell._element is cells[ci - 1]._element:
                    continue
                # 「電話」「㎡」等の固定テキストはスキップ
                t = cell.text.strip()
                if t and ("電" in t or "㎡" in t or "m2" in t):
                    continue
                if _cell_is_empty(cell):
                    _set_cell_text(cell, value)
                    wrote = True
                    break

            if wrote:
                break  # 次のフィールドへ

    _remove_seal_marks(doc)
    doc.save(output_path)
    return output_path


# ========== Excel (.xlsx) テンプレートフィラー ==========

XLSX_CONFIGS = {
    "豊島": {
        "file": "豊島区/標識設置届.xlsx",
        "sheet": "標識設置届（正）A3",
        "cells": {
            "E19": "building_name",
            "E20": "designer_info",
            "E21": "constructor_info",
            "E22": "site_address",
            "E25": "building_use",
            "I8": "applicant_address",
            "I10": "applicant_name",
            "I12": "applicant_tel",
        },
    },
    "練馬": {
        "file": "練馬区/様式（標識設置届・住民説明報告書等）.xlsx",
        "sheet": "標識設置届",
        "cells": {
            "V18": "building_name",
            "V19": "site_address",
            "V27": "structure",
            "V30": "designer_info",
            "V32": "constructor_info",
            "V35": "contact_info",
        },
    },
    "中野": {
        "file": "中野区/届出様式一式.xlsx",
        "sheet": "標識設置届",
        "cells": {
            "H23": "building_name",
            "H24": "designer_info",
            "H27": "constructor_info",
            "T8": "applicant_address",
            "T10": "applicant_name",
        },
    },
    "文京": {
        "file": "文京区/標識設置届.xlsx",
        "sheet": "表面",
        "cells": {
            "F14": "building_name",
            "F15": "designer_info",
            "F16": "constructor_info",
            "F17": "land_number",
            "F18": "zoning",
            "F19": "other_zone",
            "F20": "building_use",
            "F21": "height",
            "F22": "structure",
            "F23": "unit_count",
            "F25": "site_area",
            "F26": "building_area",
            "F27": "total_floor_area",
            "L20": "construction_type",
        },
    },
    "新宿": {
        "file": "新宿区/標識設置届.xlsx",
        "sheet": "標識設置届",
        "cells": {
            "H23": "building_name",
            "H24": "designer_info",
            "H25": "constructor_info",
            "H26": "land_number",
            "H27": "zoning",
            "H28": "other_zone",
            "H30": "building_use",
            "AB30": "construction_type",
            "H34": "structure",
        },
    },
    "葛飾": {
        "file": "葛飾区/標識設置届.xlsx",
        "sheet": "Sheet3",
        "cells": {
            "J19": "building_name",
            "J20": "site_address",
            "J21": "building_use",
            "J22": "building_area",
            "O21": "site_area",
            "O22": "total_floor_area",
            "J23": "structure",
            "O23": "foundation",
            "J26": "applicant_name",
            "J27": "designer_info",
            "J28": "constructor_info",
        },
    },
}

XLSX_REPORT_CONFIGS = {
    "港": {
        "file": "港区/隣接関係住民説明会等報告書.xlsx",
        "sheet": "Sheet1",
        "cells": {
            "F14": "building_name",
            "F15": "designer_info",
            "F18": "constructor_info",
            "F21": "site_address",
            "F24": "building_use",
        },
    },
    "練馬": {
        "file": "練馬区/様式（標識設置届・住民説明報告書等）.xlsx",
        "sheet": "住民説明報告書",
        "cells": {
            "V17": "building_name",
            "V18": "site_address",
            "V26": "structure",
            "V29": "designer_info",
            "V31": "constructor_info",
            "V33": "contact_info",
        },
    },
    "中野": {
        "file": "中野区/届出様式一式.xlsx",
        "sheet": "説明会等内容報告書(表)",
        "cells": {
            "F21": "site_address",
        },
    },
    "文京": {
        "file": "文京区/説明会等報告書.xlsx",
        "sheet": "説明会報告書（表）",
        "cells": {
            "H17": "building_name",
            "H18": "site_address",
        },
    },
}


def _fill_xlsx_with_config(config, data, output_path):
    """Excel形式テンプレートにデータを流し込む（汎用）"""
    if not openpyxl:
        return None

    template_path = os.path.join(TEMPLATES_DIR, config["file"])
    if not os.path.exists(template_path):
        return None

    wb = openpyxl.load_workbook(template_path)
    sheet_name = config.get("sheet")
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active

    fill_data = _prepare_data(data)

    for cell_ref, field_name in config["cells"].items():
        value = fill_data.get(field_name, "")
        if value:
            ws[cell_ref] = value

    # Excelシート内の「印」を削除
    for row in ws.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str) and "印" in cell.value:
                cell.value = cell.value.replace("　印", "").replace("印", "")

    wb.save(output_path)
    return output_path


# ========== メインAPI ==========

def get_available_templates(ward_name):
    """指定区で利用可能なテンプレートの種類を返す"""
    result = {"sign_notice": None, "report": None}

    if ward_name in SIGN_NOTICE_CONFIG:
        tpl = _find_template_file(ward_name, ["標識", "申請書類", "様式集"])
        if tpl:
            result["sign_notice"] = "docx"

    if ward_name in XLSX_CONFIGS:
        result["sign_notice"] = "xlsx"

    if ward_name in REPORT_CONFIG:
        tpl = _find_template_file(ward_name, ["説明", "報告"])
        if tpl:
            result["report"] = "docx"

    if ward_name in XLSX_REPORT_CONFIGS:
        result["report"] = "xlsx"

    return result


def fill_sign_notice(ward_name, data, output_path):
    """標識設置届を公式テンプレートで生成

    Returns: output_path if template was used, None if fallback needed
    """
    # Excel テンプレート
    config = XLSX_CONFIGS.get(ward_name)
    if config:
        result = _fill_xlsx_with_config(config, data, output_path)
        if result:
            return result

    # Word テンプレート
    config = SIGN_NOTICE_CONFIG.get(ward_name)
    if config:
        template = _find_template_file(ward_name, ["標識", "申請書類", "様式集"])
        if template:
            return _fill_docx_by_labels(template, config, data, output_path)

    return None


def fill_explanation_report(ward_name, data, output_path):
    """説明報告書を公式テンプレートで生成

    Returns: output_path if template was used, None if fallback needed
    """
    # Excel テンプレート
    report_config = XLSX_REPORT_CONFIGS.get(ward_name)
    if report_config:
        result = _fill_xlsx_with_config(report_config, data, output_path)
        if result:
            return result

    # Word テンプレート
    config = REPORT_CONFIG.get(ward_name)
    if config:
        template = _find_template_file(ward_name, ["説明", "報告"])
        if template:
            return _fill_docx_by_labels(template, config, data, output_path)

    return None


# ========== 解体工事テンプレート ==========

# 解体用テンプレートディレクトリ名
_DEMOLITION_SUBDIR = "demolition"


def _find_demolition_template(ward_name, filename_pattern):
    """解体用テンプレートファイルを検索"""
    ward_dir = os.path.join(TEMPLATES_DIR, f"{ward_name}区", _DEMOLITION_SUBDIR)
    if not os.path.isdir(ward_dir):
        return None
    for fname in os.listdir(ward_dir):
        if not fname.endswith(".docx"):
            continue
        if filename_pattern in fname:
            return os.path.join(ward_dir, fname)
    return None


def _set_paragraph_text(paragraph, text):
    """段落のテキストを既存書式を保持しつつ置換"""
    if paragraph.runs:
        paragraph.runs[0].text = str(text)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        run = paragraph.add_run(str(text))
        run.font.name = "游ゴシック"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
        run.font.size = Pt(10)


def _replace_checkbox(cell, mapping):
    """セル内の□/☑チェックボックスを指定マッピングに従い置換

    mapping: {"有": True, "無": False} のようなdict
    □有 → ☑有, □無 → □無 (Trueの方に☑を付ける)
    """
    for p in cell.paragraphs:
        combined = "".join(r.text for r in p.runs)
        changed = combined
        for label, checked in mapping.items():
            if checked:
                changed = changed.replace(f"□{label}", f"☑{label}")
            else:
                # 既に☑が付いていたら□に戻す
                changed = changed.replace(f"☑{label}", f"□{label}")
        if changed != combined:
            for i, run in enumerate(p.runs):
                run.text = changed if i == 0 else ""


def _get_merged_cell(table, row_idx, col_idx):
    """結合セルを考慮してセルを取得"""
    return table.cell(row_idx, col_idx)


# --- 足立区 第1号様式（解体工事のお知らせ標識）---

def _set_run_font(run, size_pt=12, font_name="ＭＳ ゴシック", bold=False):
    """runのフォントを設定"""
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.bold = bold


def _fill_adachi_demolition_sign(template_path, data, output_path):
    """足立区 解体工事のお知らせ標識（hp1gou）にデータ流し込み"""
    doc = Document(template_path)

    # .doc→.docx変換でページサイズがA4に縮小されている問題を修正
    # 標識はA3縦（29.7cm x 42.0cm）で作成する規定
    for section in doc.sections:
        section.page_width = Cm(29.7)
        section.page_height = Cm(42.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

    # ヘッダー段落を整理（P0に様式番号+注記を1行にまとめ、余分な空段落を削除）
    while len(doc.paragraphs) > 1:
        # P2以降の「A3版以上」テキストをP0に追記してから削除
        p_last = doc.paragraphs[-1]
        if p_last.text.strip():
            # P0末尾に追記
            doc.paragraphs[0].runs[-1].text += "　" + p_last.text.strip()
        p_last._element.getparent().remove(p_last._element)
        if len(doc.paragraphs) <= 1:
            break
    # P0のフォントをMS明朝12ptに統一
    for run in doc.paragraphs[0].runs:
        _set_run_font(run, 12, "ＭＳ 明朝")

    if not doc.tables:
        return None

    table = doc.tables[0]

    # --- A3に合わせて列幅・行高を調整 ---
    from docx.shared import Emu
    _col_widths = [Cm(6.3), Cm(9.2), Cm(9.7)]
    for ci, col in enumerate(table.columns):
        col.width = _col_widths[ci]
    # 各行のセル幅も設定（python-docxの制約で列幅だけでは不十分）
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = _col_widths[ci]

    _row_heights = [Cm(4.6), Cm(2.2), Cm(2.4), Cm(2.4), Cm(2.7),
                    Cm(3.7), Cm(2.1), Cm(2.3), Cm(2.3), Cm(2.3)]
    from docx.oxml.ns import qn as _qn
    for ri, row in enumerate(table.rows):
        row.height = _row_heights[ri]

    # --- Row0: タイトル「解体工事のお知らせ」を50ptに ---
    cell_title = _get_merged_cell(table, 0, 0)
    for p in cell_title.paragraphs:
        for run in p.runs:
            _set_run_font(run, 50, "ＭＳ ゴシック", bold=False)

    # --- Row6: 「お問合せ」を18pt太字に ---
    cell_inquiry = _get_merged_cell(table, 6, 0)
    for p in cell_inquiry.paragraphs:
        for run in p.runs:
            _set_run_font(run, 18, "ＭＳ ゴシック", bold=True)

    # --- データ流し込み ---
    site_address = data.get("site_address", "")
    addr_local = site_address
    for prefix in ["東京都足立区", "足立区"]:
        if addr_local.startswith(prefix):
            addr_local = addr_local[len(prefix):]
            break

    # Row1: 敷地の位置
    cell_addr = _get_merged_cell(table, 1, 1)
    _replace_in_cell(cell_addr, "足立区", f"足立区{addr_local}")

    # Row2 Col1: 高さ
    height = data.get("height", "")
    if height:
        cell_h = table.cell(2, 1)
        for p in cell_h.paragraphs:
            combined = "".join(r.text for r in p.runs)
            if "高さ" in combined and "ｍ" in combined:
                new_text = f"高さ\u3000{height}\u3000ｍ"
                for i, run in enumerate(p.runs):
                    run.text = new_text if i == 0 else ""
                break

    # Row2 Col2: 階数
    floors_above = data.get("floors_above", "")
    floors_below = data.get("floors_below", "")
    cell_fl = table.cell(2, 2)
    for p in cell_fl.paragraphs:
        combined = "".join(r.text for r in p.runs)
        if "地上" in combined and "地下" in combined:
            new_text = f"階数　地上{floors_above}階/地下{floors_below}階"
            for i, run in enumerate(p.runs):
                run.text = new_text if i == 0 else ""
            break

    # Row3 Col1: 構造
    structure = data.get("structure", "")
    if structure:
        cell_st = table.cell(3, 1)
        for p in cell_st.paragraphs:
            combined = "".join(r.text for r in p.runs)
            if "構造" in combined:
                new_text = f"構造\u3000{structure}"
                for i, run in enumerate(p.runs):
                    run.text = new_text if i == 0 else ""
                break

    # Row3 Col2: 床面積
    total_area = data.get("total_floor_area", "")
    if total_area:
        cell_area = table.cell(3, 2)
        for p in cell_area.paragraphs:
            combined = "".join(r.text for r in p.runs)
            if "床面積" in combined:
                new_text = f"床面積\u3000{total_area}\u3000㎡"
                for i, run in enumerate(p.runs):
                    run.text = new_text if i == 0 else ""
                break

    # Row4: 工事予定期間
    start_date = data.get("start_date", "")
    end_date = data.get("end_date", "")
    if start_date or end_date:
        cell_period = _get_merged_cell(table, 4, 1)
        for p in cell_period.paragraphs:
            combined = "".join(r.text for r in p.runs)
            if "年" in combined and "月" in combined:
                new_text = f"{start_date} ～ {end_date}"
                for i, run in enumerate(p.runs):
                    run.text = new_text if i == 0 else ""
                break

    # Row5: 石綿等の状況（「有・無」→選択された方に○を付ける）
    asbestos = data.get("asbestos_status", "")
    cell_asb = _get_merged_cell(table, 5, 1)
    if asbestos:
        is_yes = asbestos in ["有り", "有"]
        is_no = asbestos in ["無し", "無"]
        for p in cell_asb.paragraphs:
            combined = "".join(r.text for r in p.runs)
            if "有" in combined and "無" in combined and "・" in combined:
                if is_yes:
                    new_text = combined.replace("有\u3000・\u3000無", "〇有\u3000・\u3000無")
                elif is_no:
                    new_text = combined.replace("有\u3000・\u3000無", "有\u3000・\u3000〇無")
                else:
                    new_text = combined
                for i, run in enumerate(p.runs):
                    run.text = new_text if i == 0 else ""
                break

    # Row7: 工事施工会社（会社名/住所/電話番号）
    cell_company = _get_merged_cell(table, 7, 1)
    paragraphs = cell_company.paragraphs
    company_name = data.get("constructor_name", "")
    company_addr = data.get("constructor_address", "")
    company_tel = data.get("constructor_tel", "")
    if len(paragraphs) >= 3:
        for p in paragraphs:
            combined = "".join(r.text for r in p.runs)
            if "会社名" in combined and company_name:
                new_text = f"会社名：{company_name}"
                for i, run in enumerate(p.runs):
                    run.text = new_text if i == 0 else ""
            elif "所" in combined and "会社" not in combined and "電話" not in combined and company_addr:
                new_text = f"住\u3000\u3000所：{company_addr}"
                for i, run in enumerate(p.runs):
                    run.text = new_text if i == 0 else ""
            elif "電話番号" in combined and company_tel:
                new_text = f"電話番号：{company_tel}"
                for i, run in enumerate(p.runs):
                    run.text = new_text if i == 0 else ""

    # Row8: 現場責任者（氏名/電話）
    cell_manager = _get_merged_cell(table, 8, 1)
    manager_name = data.get("site_manager", "")
    manager_tel = data.get("constructor_tel", "")
    if manager_name:
        _replace_in_cell(cell_manager, "現場責任者氏名：", f"現場責任者氏名：{manager_name}")
    if manager_tel:
        for p in cell_manager.paragraphs:
            combined = "".join(r.text for r in p.runs)
            if "日中連絡" in combined and "電話番号" in combined:
                new_text = combined.replace("電話番号：", f"電話番号：{manager_tel}")
                for i, run in enumerate(p.runs):
                    run.text = new_text if i == 0 else ""
                break

    # Row9: 標識設置年月日（12ptに統一）
    sign_date = data.get("sign_install_date", "")
    if sign_date:
        cell_sign = _get_merged_cell(table, 9, 1)
        for p in cell_sign.paragraphs:
            combined = "".join(r.text for r in p.runs)
            if "年" in combined and "月" in combined and "日" in combined:
                # 日付プレースホルダーを実際の日付に置換
                new_text = combined
                # 「　　 年　　 月　　 日　　」部分を日付で置換
                import re as _re
                new_text = _re.sub(
                    r'[\s　]*年[\s　]*月[\s　]*日[\s　]*',
                    sign_date + "　",
                    new_text, count=1)
                for i, run in enumerate(p.runs):
                    run.text = new_text if i == 0 else ""
                    _set_run_font(run, 12, "ＭＳ ゴシック")
                break

    doc.save(output_path)
    return output_path


# --- 足立区 第2号様式（建築物解体工事事前周知報告書）---

def _fill_adachi_demolition_report(template_path, data, output_path):
    """足立区 解体工事事前周知報告書（hphoukoku）にデータ流し込み"""
    doc = Document(template_path)

    # .doc→.docx変換でマージンが大きすぎてテーブルがはみ出す問題を修正
    # テーブル幅17.2cm → 左右マージンを狭めて収める
    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # --- 段落フィールド ---
    # P3: 日付、P8: 発注者氏名、P9: 住所、P10: 電話番号
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        # 日付行（「年　　月　　日」）
        if "年" in text and "月" in text and "日" in text and i < 6:
            submit_date = data.get("submit_date", "")
            if submit_date:
                _set_paragraph_text(p, f"　　　　　　　　　　　　　　　　　　　　　　　　{submit_date}")
        # 発注者氏名
        elif "発注者等の氏名" in text:
            name = data.get("applicant_name", "")
            if name:
                _set_paragraph_text(p, f"　　　　　　　　発注者等の氏名　{name}")
        # 住所
        elif text.startswith("住") and "所" in text and i > 6:
            addr = data.get("applicant_address", "")
            if addr:
                _set_paragraph_text(p, f"　　　　　　　　住　　所　{addr}")
        # 電話番号
        elif "電話番号" in text:
            tel = data.get("applicant_tel", "")
            if tel:
                _set_paragraph_text(p, f"　　　　　　　　電話番号　{tel}")

    if not doc.tables:
        doc.save(output_path)
        return output_path

    table = doc.tables[0]
    site_address = data.get("site_address", "")
    addr_local = site_address
    for prefix in ["東京都足立区", "足立区"]:
        if addr_local.startswith(prefix):
            addr_local = addr_local[len(prefix):]
            break

    # Row0: 工事の名称
    site_name = data.get("site_name", "")
    if site_name:
        cell = _get_merged_cell(table, 0, 2)
        _set_cell_text(cell, site_name)

    # Row1: 所在地（「足立区」の後に住所）
    cell_addr = _get_merged_cell(table, 1, 2)
    _replace_in_cell(cell_addr, "足立区", f"足立区{addr_local}")

    # Row2: 工事期間
    start_date = data.get("start_date", "")
    end_date = data.get("end_date", "")
    if start_date or end_date:
        cell_period = _get_merged_cell(table, 2, 2)
        for p in cell_period.paragraphs:
            combined = "".join(r.text for r in p.runs)
            if "年" in combined:
                new_text = f"{start_date} ～ {end_date}"
                for i, run in enumerate(p.runs):
                    run.text = new_text if i == 0 else ""
                break

    # Row3: 施工者（住所/氏名/電話）
    cell_contractor = _get_merged_cell(table, 3, 2)
    c_name = data.get("constructor_name", "")
    c_addr = data.get("constructor_address", "")
    c_tel = data.get("constructor_tel", "")
    for p in cell_contractor.paragraphs:
        combined = "".join(r.text for r in p.runs)
        if "住所" in combined:
            new_text = f"住所　{c_addr}"
            for i, run in enumerate(p.runs):
                run.text = new_text if i == 0 else ""
        elif "氏名" in combined:
            new_text = f"氏名　{c_name}　電話（{c_tel}）"
            for i, run in enumerate(p.runs):
                run.text = new_text if i == 0 else ""

    # Row4: 解体建築物の概要（Col3: 床面積、Col5: 構造・階数）
    total_area = data.get("total_floor_area", "")
    if total_area:
        cell_area = table.cell(4, 3)
        _set_cell_text(cell_area, f"{total_area}㎡")
    structure = data.get("structure", "")
    floors_above = data.get("floors_above", "")
    floors_below = data.get("floors_below", "")
    struct_text = structure
    if floors_above or floors_below:
        struct_text += f"　地上{floors_above}階"
        if floors_below:
            struct_text += f"/地下{floors_below}階"
    if struct_text:
        cell_struct = table.cell(4, 5)
        _set_cell_text(cell_struct, struct_text)

    # Row5: 石綿の使用状況（□有 □無 → ☑/□）
    asbestos = data.get("asbestos_status", "")
    if asbestos:
        cell_asb = _get_merged_cell(table, 5, 2)
        _replace_checkbox(cell_asb, {
            "有": asbestos in ["有り", "有"],
            "無": asbestos in ["無し", "無"],
        })

    # Row6: 報告事項（□標識設置 □近隣の方への説明）
    # デフォルトで両方チェック
    cell_report = _get_merged_cell(table, 6, 2)
    _replace_checkbox(cell_report, {
        "標識設置": True,
        "近隣の方への説明": True,
    })

    # Row7: 標識設置年月日
    sign_date = data.get("sign_install_date", "")
    if sign_date:
        cell_sign = _get_merged_cell(table, 7, 2)
        for p in cell_sign.paragraphs:
            combined = "".join(r.text for r in p.runs)
            if "標識設置年月日" in combined:
                new_text = f"標識設置年月日　{sign_date}"
                for i, run in enumerate(p.runs):
                    run.text = new_text if i == 0 else ""
                break

    # Row8: 説明実施方法（□説明会 □戸別説明）
    method_checks = data.get("explanation_method_checks", [])
    if not method_checks:
        method = data.get("explanation_method", "")
        if "説明会" in method:
            method_checks.append("説明会")
        if "個別訪問" in method or "戸別" in method:
            method_checks.append("戸別説明")
    cell_method = _get_merged_cell(table, 8, 2)
    _replace_checkbox(cell_method, {
        "説明会": "説明会" in method_checks,
        "戸別説明": "戸別説明" in method_checks or "個別訪問" in method_checks,
    })

    # Row9: 説明時期
    explanation_date = data.get("explanation_date", "")
    if explanation_date:
        cell_explain = _get_merged_cell(table, 9, 2)
        for p in cell_explain.paragraphs:
            combined = "".join(r.text for r in p.runs)
            if "説明時期" in combined or "年" in combined:
                new_text = f"説明時期　{explanation_date}"
                for i, run in enumerate(p.runs):
                    run.text = new_text if i == 0 else ""
                break

    _remove_seal_marks(doc)
    doc.save(output_path)
    return output_path


# --- 汎用解体テンプレートフィラー ---

def _prepare_demolition_data(data):
    """解体用のデータを準備（共通前処理）"""
    d = _prepare_data(data)
    d["construction_year"] = data.get("construction_year", "")
    d["renovation_history"] = data.get("renovation_history", "")
    d["demolition_method"] = data.get("demolition_method", "")
    d["asbestos_status"] = data.get("asbestos_status", "")
    d["asbestos_removal_method"] = data.get("asbestos_removal_method", "")
    d["transport_route"] = data.get("transport_route", "")
    d["vehicle_route"] = data.get("vehicle_route", "")
    d["constructor_address"] = data.get("constructor_address", "")
    d["subcontractor_name"] = data.get("subcontractor_name", "")
    d["subcontractor_address"] = data.get("subcontractor_address", "")
    d["subcontractor_tel"] = data.get("subcontractor_tel", "")
    d["asbestos_level"] = data.get("asbestos_level", "")
    d["asbestos_locations"] = data.get("asbestos_locations", "")
    d["asbestos_types"] = data.get("asbestos_types", "")
    d["asbestos_survey_date"] = data.get("asbestos_survey_date", "")
    d["asbestos_survey_company"] = data.get("asbestos_survey_company", "")
    d["asbestos_surveyor"] = data.get("asbestos_surveyor", "")
    d["asbestos_area"] = data.get("asbestos_area", "")
    if d["start_date"] or d["end_date"]:
        d["period_text"] = f"{d['start_date']} ～ {d['end_date']}"
    else:
        d["period_text"] = ""
    d["applicant_info"] = d["applicant_name"]
    if d["applicant_address"]:
        d["applicant_info"] = f"{d['applicant_address']}　{d['applicant_name']}"
    parts = [p for p in [d["constructor_address"], d["constructor_name"]] if p]
    if d["constructor_tel"]:
        parts.append(f"TEL: {d['constructor_tel']}")
    d["constructor_full"] = "　".join(parts)
    return d


def _fill_demolition_docx_by_labels(template_path, label_map, data, output_path,
                                     table_index=0, page_size=None, margins=None):
    """Word解体テンプレートにラベル検索方式でデータを流し込む（汎用）"""
    doc = Document(template_path)
    fill_data = _prepare_demolition_data(data)
    if page_size:
        for section in doc.sections:
            section.page_width = Cm(page_size[0])
            section.page_height = Cm(page_size[1])
    if margins:
        for section in doc.sections:
            if "left" in margins:
                section.left_margin = Cm(margins["left"])
            if "right" in margins:
                section.right_margin = Cm(margins["right"])
            if "top" in margins:
                section.top_margin = Cm(margins["top"])
            if "bottom" in margins:
                section.bottom_margin = Cm(margins["bottom"])
    table = doc.tables[table_index] if table_index < len(doc.tables) else None
    if not table:
        doc.save(output_path)
        return output_path
    for label_pattern, field_name in label_map:
        value = fill_data.get(field_name, "")
        if not value:
            continue
        for ri, row in enumerate(table.rows):
            found = False
            for cell in row.cells:
                if re.search(label_pattern, cell.text):
                    found = True
                    break
            if not found:
                continue
            cells = list(row.cells)
            label_cells = set()
            for ci, cell in enumerate(cells):
                if re.search(label_pattern, cell.text):
                    label_cells.add(ci)
            wrote = False
            for ci, cell in enumerate(cells):
                if ci in label_cells:
                    continue
                if ci > 0 and cell._element is cells[ci - 1]._element:
                    continue
                t = cell.text.strip()
                if t and ("電" in t or "㎡" in t):
                    continue
                if _cell_is_empty(cell):
                    _set_cell_text(cell, value)
                    wrote = True
                    break
            if wrote:
                break
    # 段落内の日付・届出者置換
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if "年" in text and "月" in text and "日" in text and i < 6:
            sd = fill_data.get("submit_date", "")
            if sd and not any(c.isdigit() for c in text):
                _set_paragraph_text(p, f"\u3000" * 24 + sd)
        elif "発注者" in text and "氏名" in text:
            name = fill_data.get("applicant_name", "")
            if name:
                _set_paragraph_text(p, text.replace("氏名", f"氏名　{name}"))
    _remove_seal_marks(doc)
    doc.save(output_path)
    return output_path


def _xlsx_find_master_cell(ws, cell_ref):
    """結合セルの場合、書き込み可能なマスターセルを返す"""
    from openpyxl.cell.cell import MergedCell
    cell = ws[cell_ref]
    if not isinstance(cell, MergedCell):
        return cell
    # 結合範囲からマスターセルを探す
    for merged_range in ws.merged_cells.ranges:
        if cell.coordinate in merged_range:
            return ws.cell(row=merged_range.min_row, column=merged_range.min_col)
    return cell


def _fill_demolition_xlsx(template_path, cell_map, data, output_path, sheet_name=None):
    """Excel解体テンプレートにセル座標方式でデータを流し込む（汎用）"""
    if not openpyxl:
        return None
    if not os.path.exists(template_path):
        return None
    wb = openpyxl.load_workbook(template_path)
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
    fill_data = _prepare_demolition_data(data)
    for cell_ref, field_name in cell_map.items():
        value = fill_data.get(field_name, "")
        if value:
            cell = _xlsx_find_master_cell(ws, cell_ref)
            cell.value = value
    for row in ws.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str) and "印" in cell.value:
                cell.value = cell.value.replace("　印", "").replace("印", "")
    wb.save(output_path)
    return output_path


def _find_demolition_template_xlsx(ward_name, filename_pattern):
    """解体用Excelテンプレートファイルを検索"""
    ward_dir = os.path.join(TEMPLATES_DIR, f"{ward_name}区", _DEMOLITION_SUBDIR)
    if not os.path.isdir(ward_dir):
        return None
    for fname in os.listdir(ward_dir):
        if not fname.endswith(".xlsx"):
            continue
        if filename_pattern in fname:
            return os.path.join(ward_dir, fname)
    return None


# --- 解体テンプレート メインAPI ---

DEMOLITION_TEMPLATE_CONFIGS = {
    "足立": {
        "sign": {"filename": "hp1gou_3", "filler": "custom", "func": "_fill_adachi_demolition_sign"},
        "report": {"filename": "hphoukoku_4", "filler": "custom", "func": "_fill_adachi_demolition_report"},
    },
    "新宿": {
        "sign": {"filename": "標識_第1号様式", "filler": "docx", "labels": [
            (r"工事の名称|工事名称", "site_name"), (r"工\s*期|工事予定期間", "period_text"),
            (r"標識設置年月日", "sign_install_date")]},
        "report": {"filename": "報告書_第2号様式", "filler": "docx", "labels": [
            (r"工事の名称", "site_name"), (r"所\s*在\s*地|敷地", "site_address"),
            (r"標識設置年月日", "sign_install_date"), (r"説明.*時期|説明.*日", "explanation_date")]},
    },
    "文京": {
        "sign": {"filename": "標識", "filler": "docx", "labels": [
            (r"解体工事の名称|工事の名称", "site_name"), (r"工事期間|工\s*期", "period_text"),
            (r"標識設置年月日", "sign_install_date")]},
        "report": {"filename": "報告書", "filler": "docx", "labels": [
            (r"工事の名称", "site_name"), (r"所\s*在\s*地|敷地", "site_address"),
            (r"施工者", "constructor_name"), (r"現場責任者", "site_manager"),
            (r"標識設置年月日", "sign_install_date"), (r"説明.*時期|説明.*日", "explanation_date")]},
    },
    "台東": {
        "sign": {"filename": "標識_様式1", "filler": "docx", "labels": [
            (r"解体工事の名称|工事の名称", "site_name"), (r"敷地.*位置|所\s*在\s*地", "site_address"),
            (r"解体工事期間|工\s*期", "period_text"), (r"標識設置年月日", "sign_install_date")]},
        "report": {"filename": "報告書_様式2", "filler": "docx", "labels": [
            (r"工事の名称", "site_name"), (r"所\s*在\s*地|敷地", "site_address"),
            (r"施工者", "constructor_name"), (r"標識設置年月日", "sign_install_date"),
            (r"説明.*時期|説明.*日", "explanation_date")]},
    },
    "墨田": {
        "sign": {"filename": "標識_第1号様式", "filler": "docx", "labels": [
            (r"解体.*名称|工事の名称", "site_name"), (r"敷地.*位置|所\s*在\s*地", "site_address"),
            (r"元請.*業者|工事施工者", "constructor_full"), (r"解体工事期間|工\s*期", "period_text"),
            (r"標識設置年月日", "sign_install_date")]},
        "report": {"filename": "報告書_第2号様式", "filler": "docx", "labels": [
            (r"工事の名称", "site_name"), (r"所\s*在\s*地|敷地", "site_address"),
            (r"標識設置年月日", "sign_install_date"), (r"説明.*時期|説明.*日", "explanation_date")]},
    },
    "江東": {
        "sign": {"filename": "標識_第1号様式", "filler": "docx", "labels": [
            (r"名\s*称", "site_name"), (r"所\s*在\s*地", "site_address"),
            (r"解体等工事期間|工\s*期", "period_text"), (r"標識設置年月日", "sign_install_date")]},
        "report": {"filename": "報告書_第2号様式", "filler": "docx", "labels": [
            (r"工事の名称", "site_name"), (r"所\s*在\s*地|敷地", "site_address"),
            (r"施工者", "constructor_name"), (r"標識設置年月日", "sign_install_date"),
            (r"説明.*時期|説明.*日", "explanation_date")]},
    },
    "品川": {
        "sign": {"filename": "標識_第1号様式", "filler": "docx", "labels": [
            (r"敷地.*所在地|所\s*在\s*地", "site_address"), (r"事業主|発注者", "applicant_info"),
            (r"工\s*期|工事予定期間", "period_text"), (r"標識設置年月日", "sign_install_date")]},
        "report": {"filename": "説明会報告_第5号様式", "filler": "docx", "labels": [
            (r"工事の名称", "site_name"), (r"所\s*在\s*地|敷地", "site_address"),
            (r"標識設置年月日", "sign_install_date")]},
    },
    "大田": {
        "sign": {"filename": "標識_第1号様式", "filler": "docx", "labels": [
            (r"工事の場所|所\s*在\s*地", "site_address"), (r"事業主|発注者", "applicant_info"),
            (r"工\s*期|工事予定期間", "period_text"), (r"標識設置年月日", "sign_install_date")]},
        "report": {"filename": "報告書_第2号様式", "filler": "docx", "labels": [
            (r"工事の名称", "site_name"), (r"所\s*在\s*地|敷地", "site_address"),
            (r"施工者", "constructor_name"), (r"標識設置年月日", "sign_install_date"),
            (r"説明.*時期|説明.*日", "explanation_date")]},
    },
    "世田谷": {
        "sign": {"filename": "第1_2_3号様式", "filler": "docx", "table_index": 0, "labels": [
            (r"工事の名称", "site_name"), (r"敷地.*位置|所\s*在\s*地", "site_address"),
            (r"事業主|発注者", "applicant_info"), (r"施工者|元請", "constructor_full"),
            (r"工\s*期|工事予定期間", "period_text"), (r"標識設置年月日", "sign_install_date")]},
    },
    "渋谷": {
        "sign": {"filename": "標識_第2号様式", "filler": "docx", "labels": [
            (r"解体工事の名称|工事の名称", "site_name"), (r"敷地.*位置|所\s*在\s*地", "site_address"),
            (r"解体工事の工期|工\s*期", "period_text"), (r"標識設置年月日", "sign_install_date")]},
        "report": {"filename": "説明会報告書_第5号様式", "filler": "docx", "labels": [
            (r"工事の名称", "site_name"), (r"所\s*在\s*地|敷地", "site_address"),
            (r"標識設置年月日", "sign_install_date")]},
    },
    "中野": {
        "sign": {"filename": "標識", "filler": "docx", "labels": [
            (r"解体工事の名称|工事の名称", "site_name"), (r"敷地.*位置|所\s*在\s*地", "site_address"),
            (r"工事施工者|施工者", "constructor_full"), (r"工\s*期|工事予定期間", "period_text"),
            (r"標識設置年月日", "sign_install_date")]},
    },
    "北": {
        "sign": {"filename": "標識_第1号様式", "filler": "docx", "labels": [
            (r"敷地.*位置|住居表示", "site_address"), (r"事業主|発注者", "applicant_info"),
            (r"工事予定期間|工\s*期", "period_text"), (r"標識設置年月日", "sign_install_date")]},
        "report": {"filename": "説明会報告書_第3号様式", "filler": "docx", "labels": [
            (r"工事の名称", "site_name"), (r"所\s*在\s*地|敷地", "site_address"),
            (r"標識設置年月日", "sign_install_date"), (r"説明.*時期|説明.*日", "explanation_date")]},
    },
    "荒川": {
        "sign": {"filename": "標識", "filler": "docx", "labels": [
            (r"解体工事の名称|工事の名称", "site_name"), (r"敷地.*位置|所\s*在\s*地", "site_address"),
            (r"工事予定期間|工\s*期", "period_text"), (r"標識設置年月日", "sign_install_date")]},
        "report": {"filename": "報告書", "filler": "docx", "labels": [
            (r"工事の名称", "site_name"), (r"所\s*在\s*地|敷地", "site_address"),
            (r"標識設置年月日", "sign_install_date")]},
    },
    "千代田": {
        "sign": {"filename": "標識_第1号様式", "filler": "xlsx", "sheet": "◎標識", "cells": {
            "E7": "site_name", "E8": "site_address", "E11": "applicant_address",
            "E12": "applicant_name", "E13": "constructor_address", "E14": "constructor_name",
            "E19": "asbestos_survey_date", "D28": "constructor_address",
            "D29": "site_manager", "D30": "constructor_tel"}},
        "report": {"filename": "報告書_第2号様式", "filler": "docx", "labels": [
            (r"工事の名称", "site_name"), (r"所\s*在\s*地|敷地", "site_address"),
            (r"施工者", "constructor_name"), (r"標識設置年月日", "sign_install_date"),
            (r"説明.*時期|説明.*日", "explanation_date")]},
    },
    "中央": {
        "sign": {"filename": "標識_第1号様式", "filler": "xlsx", "sheet": "１号様式", "cells": {
            "E7": "site_name", "E8": "applicant_address", "E9": "applicant_name",
            "E10": "constructor_address", "E11": "constructor_name", "H12": "site_address",
            "E28": "constructor_address", "E29": "site_manager"}},
        "report": {"filename": "説明会報告_第3号様式", "filler": "xlsx", "sheet": "３号様式 ", "cells": {
            "N8": "applicant_address", "N10": "applicant_name",
            "E12": "site_name", "H13": "site_address"}},
    },
    "港": {
        "sign": {"filename": "標識_第2号様式", "filler": "xlsx", "sheet": "Sheet1", "cells": {
            "D9": "site_name", "D13": "applicant_address", "D15": "applicant_name",
            "D27": "constructor_address", "D28": "site_manager", "D29": "constructor_tel"}},
        "report": {"filename": "説明会報告書_第4号様式", "filler": "xlsx", "sheet": "Sheet1", "cells": {
            "C14": "site_name", "C18": "constructor_address", "C19": "constructor_name"}},
    },
    "杉並": {
        "sign": {"filename": "標識", "filler": "xlsx", "sheet": "工事看板", "cells": {
            "E5": "site_name", "K5": "site_address", "K11": "constructor_name",
            "K14": "constructor_address", "K20": "constructor_tel",
            "K22": "site_manager", "K23": "constructor_tel"}},
    },
    "豊島": {
        "sign": {"filename": "標識_お知らせ", "filler": "xlsx", "cells": {
            "E5": "site_name", "E6": "site_address"}},
    },
    "葛飾": {
        "sign": {"filename": "標識_お知らせ", "filler": "xlsx",
                 "sheet": "石綿未使用 (ひな形) ", "cells": {
            "E9": "site_name", "E14": "constructor_name", "E15": "constructor_address"}},
        "report": {"filename": "報告書様式", "filler": "xlsx", "cells": {
            "E5": "site_name", "E6": "site_address"}},
    },
    "練馬": {
        "sign": {"filename": "標識様式例", "filler": "xlsx", "cells": {
            "D5": "site_name", "D6": "site_address"}},
        "report": {"filename": "住民説明実施報告書", "filler": "docx", "labels": [
            (r"工事の名称", "site_name"), (r"所\s*在\s*地|敷地", "site_address"),
            (r"施工者", "constructor_name"), (r"標識設置年月日", "sign_install_date"),
            (r"説明.*時期|説明.*日", "explanation_date")]},
    },
    "江戸川": {
        "sign": {"filename": "標識_アスベストなし", "filler": "xlsx",
                 "sheet": "レベル3 届出不要・石綿なし(白紙)", "cells": {
            "E6": "site_name", "F8": "constructor_name", "F11": "constructor_address",
            "F14": "site_manager", "F15": "constructor_tel"}},
    },
    "目黒": {
        "sign": {"filename": "標識設置届_作成補助ツール", "filler": "xlsx", "cells": {
            "D5": "site_name", "D6": "site_address"}},
    },
}

_DEMOLITION_FILLERS = {
    "_fill_adachi_demolition_sign": _fill_adachi_demolition_sign,
    "_fill_adachi_demolition_report": _fill_adachi_demolition_report,
}


def get_available_demolition_templates(ward_name):
    """指定区で利用可能な解体テンプレートの種類を返す"""
    result = {"sign": None, "report": None}
    config = DEMOLITION_TEMPLATE_CONFIGS.get(ward_name)
    if not config:
        return result
    for doc_type in ("sign", "report"):
        cfg = config.get(doc_type)
        if not cfg:
            continue
        ft = cfg.get("filler", "")
        if ft == "docx" or ft == "custom":
            tpl = _find_demolition_template(ward_name, cfg["filename"])
            if tpl:
                result[doc_type] = "docx"
        elif ft == "xlsx":
            tpl = _find_demolition_template_xlsx(ward_name, cfg["filename"])
            if tpl:
                result[doc_type] = "xlsx"
    return result


def fill_demolition_sign(ward_name, data, output_path):
    """解体工事のお知らせ標識を公式テンプレートで生成"""
    config = DEMOLITION_TEMPLATE_CONFIGS.get(ward_name)
    if not config or "sign" not in config:
        return None
    cfg = config["sign"]
    ft = cfg.get("filler", "")
    if ft == "custom":
        tpl = _find_demolition_template(ward_name, cfg["filename"])
        if not tpl:
            return None
        func = _DEMOLITION_FILLERS.get(cfg.get("func"))
        return func(tpl, data, output_path) if func else None
    elif ft == "docx":
        tpl = _find_demolition_template(ward_name, cfg["filename"])
        if not tpl:
            return None
        return _fill_demolition_docx_by_labels(
            tpl, cfg["labels"], data, output_path,
            table_index=cfg.get("table_index", 0),
            page_size=cfg.get("page_size"), margins=cfg.get("margins"))
    elif ft == "xlsx":
        tpl = _find_demolition_template_xlsx(ward_name, cfg["filename"])
        if not tpl:
            return None
        if output_path.endswith(".docx"):
            output_path = output_path.replace(".docx", ".xlsx")
        return _fill_demolition_xlsx(tpl, cfg["cells"], data, output_path, sheet_name=cfg.get("sheet"))
    return None


def fill_demolition_report(ward_name, data, output_path):
    """解体工事事前周知報告書を公式テンプレートで生成"""
    config = DEMOLITION_TEMPLATE_CONFIGS.get(ward_name)
    if not config or "report" not in config:
        return None
    cfg = config["report"]
    ft = cfg.get("filler", "")
    if ft == "custom":
        tpl = _find_demolition_template(ward_name, cfg["filename"])
        if not tpl:
            return None
        func = _DEMOLITION_FILLERS.get(cfg.get("func"))
        return func(tpl, data, output_path) if func else None
    elif ft == "docx":
        tpl = _find_demolition_template(ward_name, cfg["filename"])
        if not tpl:
            return None
        return _fill_demolition_docx_by_labels(
            tpl, cfg["labels"], data, output_path,
            table_index=cfg.get("table_index", 0), margins=cfg.get("margins"))
    elif ft == "xlsx":
        tpl = _find_demolition_template_xlsx(ward_name, cfg["filename"])
        if not tpl:
            return None
        if output_path.endswith(".docx"):
            output_path = output_path.replace(".docx", ".xlsx")
        return _fill_demolition_xlsx(tpl, cfg["cells"], data, output_path, sheet_name=cfg.get("sheet"))
    return None


# ========== 区ごとの必要フィールド定義 ==========

# フィールドID → 表示名のマスター定義
FIELD_LABELS = {
    "building_name": "建築物の名称",
    "site_address": "建築場所（住居表示）",
    "land_number": "地名地番",
    "building_use": "主要用途",
    "structure": "構造",
    "foundation": "基礎工法",
    "construction_type": "工事種別",
    "height": "高さ（m）",
    "floors_above": "地上階数",
    "floors_below": "地下階数",
    "unit_count": "総住戸数",
    "oneroom_count": "ワンルーム戸数（40㎡未満）",
    "site_area": "敷地面積（㎡）",
    "building_area": "建築面積（㎡）",
    "total_floor_area": "延べ面積（㎡）",
    "zoning": "用途地域",
    "fire_zone": "防火地域",
    "other_zone": "その他の地域・地区",
    "start_date": "着工予定日",
    "end_date": "完了予定日",
    "submit_date": "届出日",
    "sign_install_date": "標識設置日",
    "applicant_name": "届出者（建築主）氏名",
    "applicant_address": "届出者 住所",
    "applicant_tel": "届出者 電話",
    "designer_name": "設計者名",
    "designer_tel": "設計者 電話",
    "constructor_name": "施工者名",
    "constructor_tel": "施工者 電話",
    "site_manager": "現場責任者",
    "explanation_date": "説明実施日",
    "explanation_method": "説明方法",
    "target_count": "説明対象戸数",
    "explained_count": "説明済み戸数",
    "unexplained_count": "未説明戸数",
    "opinions": "住民からの意見・要望",
}


def get_required_fields(ward_name):
    """指定区のテンプレートが必要とするフィールドIDのセットを返す

    Returns:
        dict: {"sign_notice": set of field IDs, "report": set of field IDs}
    """
    sign_fields = set()
    report_fields = set()

    # ラベル→フィールドIDのマッピングから抽出（info系は元フィールドに展開）
    _info_expand = {
        "designer_info": {"designer_name", "designer_tel"},
        "constructor_info": {"constructor_name", "constructor_tel"},
        "contact_info": {"site_manager", "constructor_tel"},
        "floors_text": {"floors_above", "floors_below"},
    }

    def _expand(field_id):
        return _info_expand.get(field_id, {field_id})

    # Word標識設置届
    config = SIGN_NOTICE_CONFIG.get(ward_name)
    if config:
        for _, field_id in config["labels"]:
            sign_fields |= _expand(field_id)

    # Excel標識設置届
    config = XLSX_CONFIGS.get(ward_name)
    if config:
        for _, field_id in config["cells"].items():
            sign_fields |= _expand(field_id)

    # Word報告書
    config = REPORT_CONFIG.get(ward_name)
    if config:
        for _, field_id in config["labels"]:
            report_fields |= _expand(field_id)

    # Excel報告書
    config = XLSX_REPORT_CONFIGS.get(ward_name)
    if config:
        for _, field_id in config["cells"].items():
            report_fields |= _expand(field_id)

    # テンプレートがない場合（汎用生成）でも最低限のフィールドは必要
    if not sign_fields:
        sign_fields = {
            "building_name", "site_address", "building_use", "structure",
            "height", "floors_above", "floors_below",
            "site_area", "building_area", "total_floor_area",
            "start_date", "end_date", "submit_date", "sign_install_date",
            "applicant_name", "applicant_address", "applicant_tel",
            "designer_name", "designer_tel",
            "constructor_name", "constructor_tel",
        }
    if not report_fields:
        report_fields = {
            "building_name", "site_address", "building_use", "structure",
            "height", "floors_above", "floors_below",
            "start_date", "end_date", "submit_date",
            "applicant_name", "applicant_address", "applicant_tel",
            "explanation_date", "explanation_method",
            "target_count", "explained_count", "unexplained_count",
            "opinions",
        }

    return {"sign_notice": sign_fields, "report": report_fields}
