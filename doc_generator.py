# -*- coding: utf-8 -*-
"""近隣説明会 届出書類 Word生成モジュール
東京都の様式をベースに以下を生成:
  1. 標識設置届（建築計画のお知らせ）
  2. 近隣説明報告書
  3. 工事のお知らせ（住民配布用チラシ）
  4. 近隣説明範囲図（地図付きWord）
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from ward_config import get_ward_config, get_demolition_checkboxes, DEMOLITION_CHECKBOX_DEFS
from template_filler import fill_sign_notice as _fill_official_sign_notice
from template_filler import fill_explanation_report as _fill_official_report


# ========== ユーティリティ ==========

def _set_cell(cell, text, font_size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    """セルにテキストを設定"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = "游ゴシック"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
    run.bold = bold


def _add_heading_paragraph(doc, text, font_size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER):
    """見出し段落を追加"""
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "游ゴシック"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
    run.bold = bold
    return p


def _add_body_paragraph(doc, text, font_size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    """本文段落を追加"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "游ゴシック"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
    run.bold = bold
    return p


def _render_checks(options, selected, separator="　　"):
    """チェックボックスリストを ☑/□ テキストに変換

    Args:
        options: 選択肢のリスト ["有り", "無し", "調査中"]
        selected: 選択された値（str or list）
        separator: 選択肢間の区切り文字
    Returns:
        str: "☑ 有り　　□ 無し　　□ 調査中"
    """
    if isinstance(selected, str):
        selected = [selected]
    parts = []
    for opt in options:
        mark = "☑" if opt in selected else "□"
        parts.append(f"{mark} {opt}")
    return separator.join(parts)


def _set_cell_with_checks(cell, options, selected, font_size=10, separator="　　"):
    """セルにチェックボックス付きテキストを設定"""
    text = _render_checks(options, selected, separator)
    _set_cell(cell, text, font_size=font_size)


def _set_table_borders(table):
    """テーブルに罫線を設定"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.makeelement(
            qn(f"w:{edge}"),
            {qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "0", qn("w:color"): "000000"},
        )
        borders.append(element)
    tblPr.append(borders)


# ========== 1. 標識設置届 ==========

def generate_sign_notice(data, output_path):
    """標識設置届（建築計画のお知らせ）を生成
    公式テンプレートがある区はそちらを使用、なければ自作生成
    """
    ward_name = data.get("ward_name", "")

    # 公式テンプレートがあればそちらを使用
    official = _fill_official_sign_notice(ward_name, data, output_path)
    if official:
        return official

    # フォールバック: 自作Word生成
    wc = get_ward_config(ward_name)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ヘッダー
    _add_heading_paragraph(doc, "標 識 設 置 届", font_size=18)
    _add_body_paragraph(doc, f"（{wc['ordinance_name']}{wc['sign_article']}の規定による届出）",
                        font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # 日付・宛先
    _add_body_paragraph(doc, f"　　　　　　　　　　　　　　　　　　　　　　　　{data.get('submit_date', '令和　年　月　日')}")
    _add_body_paragraph(doc, f"　{ward_name}{wc['suffix']}　殿")
    _add_body_paragraph(doc, "")

    # 届出者情報（右寄せ）
    _add_body_paragraph(doc, f"届出者　住所　{data.get('applicant_address', '')}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_body_paragraph(doc, f"氏名　{data.get('applicant_name', '')}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_body_paragraph(doc, f"電話　{data.get('applicant_tel', '')}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_body_paragraph(doc, "")
    _add_body_paragraph(doc, "　下記のとおり標識を設置しましたので届け出ます。", space_after=12)

    # メインテーブル
    rows_data = [
        ("建築物の名称", data.get("building_name", "")),
        ("建築場所", data.get("site_address", "")),
        ("用途", data.get("building_use", "")),
        ("敷地面積", data.get("site_area", "") + " ㎡"),
        ("建築面積", data.get("building_area", "") + " ㎡"),
        ("延べ面積", data.get("total_floor_area", "") + " ㎡"),
        ("構造", data.get("structure", "")),
        ("階数", f"地上 {data.get('floors_above', '')} 階　地下 {data.get('floors_below', '')} 階"),
        ("高さ", data.get("height", "") + " m"),
        ("着工予定日", data.get("start_date", "")),
        ("完了予定日", data.get("end_date", "")),
        ("設計者", f"{data.get('designer_name', '')}　TEL: {data.get('designer_tel', '')}"),
        ("施工者", f"{data.get('constructor_name', '')}　TEL: {data.get('constructor_tel', '')}"),
        ("標識設置年月日", data.get("sign_install_date", "")),
        ("標識設置場所", data.get("sign_location", "建築予定地の道路に面する見やすい場所")),
    ]

    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    for i, (label, value) in enumerate(rows_data):
        _set_cell(table.cell(i, 0), label, font_size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(table.cell(i, 1), value, font_size=10)
        table.cell(i, 0).width = Cm(4.0)
        table.cell(i, 1).width = Cm(12.0)

    _add_body_paragraph(doc, "")
    _add_body_paragraph(doc, "備考：本届出書には、案内図及び配置図を添付すること。", font_size=9)

    doc.save(output_path)
    return output_path


# ========== 1a. 解体工事のお知らせ標識（足立区 第1号様式準拠） ==========

def generate_demolition_sign(data, output_path):
    """解体工事のお知らせ標識を生成（足立区 別記第1号様式準拠）
    A3版以上の看板に記載する内容をWord文書として出力する。
    """
    ward_name = data.get("ward_name", "")
    wc = get_ward_config(ward_name)
    demo_cfg = wc.get("demolition", {})

    doc = Document()
    section = doc.sections[0]
    # A3横を想定（ただしWordとしてはA4縦で出力し、印刷時に拡大を案内）
    section.page_width = Cm(29.7)
    section.page_height = Cm(42.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 様式番号
    _add_body_paragraph(doc, f"別記第1号様式（{demo_cfg.get('ordinance_name', '建築物の解体工事の事前周知に関する要綱')}第7条関係）",
                        font_size=8, space_after=6)

    # タイトル
    _add_heading_paragraph(doc, "解 体 工 事 の お 知 ら せ", font_size=22)
    _add_body_paragraph(doc, "", space_after=8)

    # --- 石綿のチェック状態 ---
    _asb_raw = data.get("asbestos_status", "調査中")
    _asb_checks = _render_checks(["有り", "無し", "調査中"], _asb_raw)

    # --- 大規模建築物チェック ---
    _lb_checks_raw = data.get("large_building_checks", [])
    if isinstance(_lb_checks_raw, str):
        _lb_checks_raw = [x.strip() for x in _lb_checks_raw.split(",") if x.strip()]
    _lb_text = _render_checks(
        ["木造以外で3階以上", "地階を有する", "延べ面積500m²以上"],
        _lb_checks_raw, separator="\n")

    # メインテーブル
    rows_data = [
        ("解体建築物の所在地", data.get("site_address", "")),
        ("解体建築物の規模", f"延べ面積 {data.get('total_floor_area', '')} ㎡　　"
                          f"地上 {data.get('floors_above', '')} 階　地下 {data.get('floors_below', '')} 階"),
        ("構　　造", data.get("structure", "")),
        ("高　　さ", f"{data.get('height', '')} m"),
        ("工 事 期 間", f"{data.get('start_date', '')} から {data.get('end_date', '')} まで"),
        ("解 体 方 法", data.get("demolition_method", "")),
        ("作 業 時 間", data.get("work_hours", "午前8時00分 ～ 午後5時00分")),
        ("発注者　氏　名", data.get("applicant_name", "")),
        ("　　　　住　所", data.get("applicant_address", "")),
        ("　　　　連絡先", data.get("applicant_tel", "")),
        ("工事業者　氏　名", data.get("constructor_name", "")),
        ("　　　　　住　所", data.get("constructor_address", "")),
        ("　　　　　連絡先", data.get("constructor_tel", "")),
        ("石綿等の使用の有無", _asb_checks),
        ("石綿等の除去方法", data.get("asbestos_removal_method", "")),
        ("安全対策・公害防止対策", data.get("safety_measures",
            "・仮囲い・防音パネルの設置\n・散水による粉塵防止\n・交通誘導員の配置")),
        ("搬出経路", data.get("transport_route", "")),
        ("工事車両通行経路", data.get("vehicle_route", "")),
    ]

    # 大規模建築物チェックが区で必要な場合は行を追加
    if "large_building" in demo_cfg.get("checkbox_groups", []):
        rows_data.insert(4, ("大規模建築物等", _lb_text))

    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    for i, (label, value) in enumerate(rows_data):
        _set_cell(table.cell(i, 0), label, font_size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(table.cell(i, 1), value, font_size=12)
        table.cell(i, 0).width = Cm(6.0)
        table.cell(i, 1).width = Cm(18.0)

    _add_body_paragraph(doc, "", space_after=8)

    # 注意書き
    _add_body_paragraph(doc, "※ この標識は、足立区建築物の解体工事の事前周知に関する要綱第7条の規定に基づき設置するものです。",
                        font_size=9, space_after=4)
    _add_body_paragraph(doc, "※ A3版以上の大きさで作成し、建築敷地の道路に接する部分に、"
                        "地面から標識の下端までおおむね1メートルの高さに設置してください。",
                        font_size=9, space_after=4)
    _add_body_paragraph(doc, f"問い合わせ先: {demo_cfg.get('submit_to', '建築審査課')}",
                        font_size=9, space_after=2)

    doc.save(output_path)
    return output_path


# ========== 1b. 解体工事事前周知報告書 ==========

def _merge_cells_and_set(table, row1, col1, row2, col2, text, font_size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    """セルを結合してテキストを設定"""
    cell = table.cell(row1, col1).merge(table.cell(row2, col2))
    _set_cell(cell, text, font_size=font_size, bold=bold, align=align)
    return cell


def generate_demolition_report(data, output_path):
    """解体工事 事前周知報告書を生成（実際の区様式に準拠）"""
    ward_name = data.get("ward_name", "")
    wc = get_ward_config(ward_name)
    suffix = wc.get("suffix", "区長").replace("長", "")  # "区長"→"区", "市長"→"市"

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # --- 様式番号 ---
    _add_body_paragraph(doc, "様式（事前周知報告書）", font_size=9, space_after=2)

    # --- 日付（右寄せ） ---
    _add_body_paragraph(doc, data.get("submit_date", "　　年　　月　　日"),
                        font_size=10.5, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)

    # --- タイトル ---
    _add_heading_paragraph(doc, "事 前 周 知 報 告 書", font_size=18)
    _add_body_paragraph(doc, "", space_after=2)

    # --- 宛先 ---
    _add_body_paragraph(doc, f"　{ward_name}{wc['suffix']}　殿", font_size=11, space_after=10)

    # --- 届出者（右寄せ） ---
    _add_body_paragraph(doc, f"住　所　{data.get('applicant_address', '')}",
                        font_size=10.5, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    _add_body_paragraph(doc, f"氏　名　{data.get('applicant_name', '')}",
                        font_size=10.5, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    _add_body_paragraph(doc, "（法人にあっては名称、代表者の氏名）",
                        font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=8)

    # --- 要綱文 ---
    _demo_cfg = wc.get("demolition", {})
    _ordinance_ref = _demo_cfg.get("ordinance_name", "")
    if not _ordinance_ref:
        _ordinance_ref = f"{ward_name}{suffix}建築物等の解体等工事に係る計画の事前周知に関する要綱"
    _add_body_paragraph(doc,
        f"　{_ordinance_ref}に基づく事前周知について以下のとおり報告します。",
        font_size=10.5, space_after=8)

    # ======= メインテーブル（1つの表で全項目）=======
    # 行構成:
    #  0: 解体等工事の名称
    #  1: 解体等工事の場所
    #  2-4: 解体建築物等の概要（延べ面積/階数, 構造/高さ, 竣工年/改修歴）
    #  5: 工事予定期間
    #  6-7: 発注者（氏名, 住所/電話）
    #  8-9: 元請業者（会社名(代表者), 住所/電話）
    # 10-11: 下請負人（会社名(代表者), 住所/電話）
    # 12-13: 問合せ先担当者（会社名, 氏名/電話）
    # 14-15: 近隣説明（説明時期, 実施方法）
    # 16-18: 添付書類
    COL = 6  # 6列テーブル
    ROWS = 19
    table = doc.add_table(rows=ROWS, cols=COL)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    _addr_raw = data.get("site_address", "")
    _addr_clean = _addr_raw.replace(f"東京都{ward_name}{suffix}", "").replace(f"{ward_name}{suffix}", "")

    # --- Row 0: 解体等工事の名称 ---
    _merge_cells_and_set(table, 0, 0, 0, 1, "解体等工事の名称", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _merge_cells_and_set(table, 0, 2, 0, 5, data.get("site_name", ""), 10)

    # --- Row 1: 解体等工事の場所 ---
    _merge_cells_and_set(table, 1, 0, 1, 1, "解体等工事の場所", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _merge_cells_and_set(table, 1, 2, 1, 5, f"{ward_name}{suffix}{_addr_clean}", 10)

    # --- Row 2-4: 解体建築物等の概要 ---
    _merge_cells_and_set(table, 2, 0, 4, 0, "解体建築物等\nの　概　要", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    # Row 2: 延べ面積 / 階数
    _set_cell(table.cell(2, 1), "延べ面積", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.cell(2, 2), f"{data.get('total_floor_area', '')} ㎡", 10)
    _set_cell(table.cell(2, 3), "階　数", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _merge_cells_and_set(table, 2, 4, 2, 5,
        f"地上 {data.get('floors_above', '')} 階、地下 {data.get('floors_below', '')} 階", 10)
    # Row 3: 構造 / 高さ
    _set_cell(table.cell(3, 1), "構　造", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.cell(3, 2), data.get("structure", ""), 10)
    _set_cell(table.cell(3, 3), "高　さ", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _merge_cells_and_set(table, 3, 4, 3, 5, f"{data.get('height', '')} m", 10)
    # Row 4: 竣工年 / 改修歴
    _set_cell(table.cell(4, 1), "竣工年又は\n築年数", 8, True, WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.cell(4, 2), data.get("construction_year", ""), 10)
    _set_cell(table.cell(4, 3), "増改築、\n改修歴", 8, True, WD_ALIGN_PARAGRAPH.CENTER)
    _merge_cells_and_set(table, 4, 4, 4, 5, data.get("renovation_history", "無"), 10)

    # --- Row 5: 工事予定期間 ---
    _merge_cells_and_set(table, 5, 0, 5, 1, "工事予定期間", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _merge_cells_and_set(table, 5, 2, 5, 5,
        f"{data.get('start_date', '')} から {data.get('end_date', '')} まで", 10)

    # --- Row 6-7: 発注者 ---
    _merge_cells_and_set(table, 6, 0, 7, 1, "発　注　者\n（法人にあっては名称、\n代表者の氏名）", 8, True, WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.cell(6, 2), "氏　名", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _merge_cells_and_set(table, 6, 3, 6, 5, data.get("applicant_name", ""), 10)
    _set_cell(table.cell(7, 2), "住　所", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.cell(7, 3), data.get("applicant_address", ""), 10)
    _set_cell(table.cell(7, 4), "電　話", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.cell(7, 5), data.get("applicant_tel", ""), 10)

    # --- Row 8-9: 元請業者 ---
    _merge_cells_and_set(table, 8, 0, 9, 1, "元 請 業 者", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.cell(8, 2), "会社名（代表者）", 8, True, WD_ALIGN_PARAGRAPH.CENTER)
    _merge_cells_and_set(table, 8, 3, 8, 5, data.get("constructor_name", ""), 10)
    _set_cell(table.cell(9, 2), "住　所", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.cell(9, 3), data.get("constructor_address", ""), 10)
    _set_cell(table.cell(9, 4), "電　話", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.cell(9, 5), data.get("constructor_tel", ""), 10)

    # --- Row 10-11: 下請負人 ---
    _merge_cells_and_set(table, 10, 0, 11, 1, "下 請 負 人", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.cell(10, 2), "会社名（代表者）", 8, True, WD_ALIGN_PARAGRAPH.CENTER)
    _merge_cells_and_set(table, 10, 3, 10, 5, data.get("subcontractor_name", ""), 10)
    _set_cell(table.cell(11, 2), "住　所", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.cell(11, 3), data.get("subcontractor_address", ""), 10)
    _set_cell(table.cell(11, 4), "電　話", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.cell(11, 5), data.get("subcontractor_tel", ""), 10)

    # --- Row 12-13: 問合せ先担当者 ---
    _merge_cells_and_set(table, 12, 0, 13, 1, "問合せ先担当者", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.cell(12, 2), "会社名", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _merge_cells_and_set(table, 12, 3, 12, 5, data.get("constructor_name", ""), 10)
    _set_cell(table.cell(13, 2), "氏　名", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.cell(13, 3), data.get("site_manager", ""), 10)
    _set_cell(table.cell(13, 4), "電　話", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.cell(13, 5), data.get("constructor_tel", ""), 10)

    # --- Row 14-15: 近隣説明 ---
    _merge_cells_and_set(table, 14, 0, 15, 1, "近 隣 説 明", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    # Row 14: 説明時期
    _set_cell(table.cell(14, 2), "説明時期", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _merge_cells_and_set(table, 14, 3, 14, 5, data.get("explanation_date", ""), 10)
    # Row 15: 実施方法（☑/□ チェックボックス）
    _set_cell(table.cell(15, 2), "実施方法", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _method_checks = data.get("explanation_method_checks", [])
    if not _method_checks:
        # 旧形式（テキスト）からの変換
        _method = data.get("explanation_method", "個別訪問による説明")
        if "説明会" in _method:
            _method_checks.append("説明会")
        if "個別訪問" in _method:
            _method_checks.append("個別訪問")
        if "書面" in _method or "ポスティング" in _method:
            _method_checks.append("書面配付（ポスティング）")
    _method_text = _render_checks(["説明会", "個別訪問", "書面配付（ポスティング）"], _method_checks)
    _merge_cells_and_set(table, 15, 3, 15, 5, _method_text, 10)

    # --- Row 16-18: 添付書類（☑/□ チェックボックス） ---
    _merge_cells_and_set(table, 16, 0, 18, 1, "添 付 書 類", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    _att_checks = data.get("attachment_checks", [])
    _att_opts = [
        ("案内図（説明範囲をマーキング）", "案内図（説明を行った家等が分かるようにマーキングすること）"),
        ("配布チラシの写し", "説明に使用したチラシ等（近隣説明範囲図を含む）"),
        ("標識設置写真（遠景・近景）", "工事対象建物の写真（遠景、近景等）"),
    ]
    for idx, (check_key, display_text) in enumerate(_att_opts):
        _mark = "☑" if check_key in _att_checks else "□"
        _merge_cells_and_set(table, 16 + idx, 2, 16 + idx, 5,
            f"{_mark} {display_text}", 9)

    _add_body_paragraph(doc, "", space_after=4)

    # --- 標識設置報告（足立区等で必要） ---
    if data.get("sign_install_date"):
        _add_body_paragraph(doc, "【標識設置の報告】", font_size=10, bold=True, space_after=4)
        sign_rows = [
            ("標識設置日", data.get("sign_install_date", "")),
            ("標識設置場所", data.get("sign_location", "建築敷地の道路に面する見やすい場所")),
        ]
        tbl_sign = doc.add_table(rows=len(sign_rows), cols=2)
        tbl_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(tbl_sign)
        for i, (label, value) in enumerate(sign_rows):
            _set_cell(tbl_sign.cell(i, 0), label, font_size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell(tbl_sign.cell(i, 1), value, font_size=10)
            tbl_sign.cell(i, 0).width = Cm(4.0)
            tbl_sign.cell(i, 1).width = Cm(13.0)
        _add_body_paragraph(doc, "※ 標識設置状況の写真を別紙に添付してください。", font_size=8, space_after=6)

    # --- 石綿等の調査結果（要綱で事前調査が必要） ---
    _asbestos = data.get("asbestos_status", "")
    if _asbestos:
        _add_body_paragraph(doc, "【石綿等の調査結果】", font_size=10, bold=True, space_after=4)
        _asb_display = _render_checks(["有り", "無し", "調査中"], _asbestos)
        asbestos_rows = [
            ("石綿等の使用の有無", _asb_display),
            ("石綿等の除去方法", data.get("asbestos_removal_method", "該当なし")),
        ]
        tbl_asb = doc.add_table(rows=len(asbestos_rows), cols=2)
        tbl_asb.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(tbl_asb)
        for i, (label, value) in enumerate(asbestos_rows):
            _set_cell(tbl_asb.cell(i, 0), label, font_size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell(tbl_asb.cell(i, 1), value, font_size=10)
            tbl_asb.cell(i, 0).width = Cm(4.0)
            tbl_asb.cell(i, 1).width = Cm(13.0)
        _add_body_paragraph(doc, "", space_after=4)

    # --- 大規模建築物等チェック ---
    _cb_groups = _demo_cfg.get("checkbox_groups", []) if _demo_cfg else []
    if "large_building" in _cb_groups:
        _lb_raw = data.get("large_building_checks", [])
        if isinstance(_lb_raw, str):
            _lb_raw = [x.strip() for x in _lb_raw.split(",") if x.strip()]
        _lb_def = DEMOLITION_CHECKBOX_DEFS["large_building"]
        _add_body_paragraph(doc, "【大規模建築物等の該当】", font_size=10, bold=True, space_after=4)
        _lb_display = _render_checks(_lb_def["options"], _lb_raw, separator="　　")
        _add_body_paragraph(doc, f"　{_lb_display}", font_size=10, space_after=6)

    # --- 特定建設作業チェック ---
    if "specific_construction" in _cb_groups:
        _sc_val = data.get("specific_construction_status", "")
        _sc_display = _render_checks(["該当する", "該当しない"], _sc_val)
        _add_body_paragraph(doc, "【特定建設作業（騒音・振動規制法）】", font_size=10, bold=True, space_after=4)
        _add_body_paragraph(doc, f"　{_sc_display}", font_size=10, space_after=6)

    # --- ねずみ駆除チェック ---
    if "rodent_control" in _cb_groups:
        _rc_val = data.get("rodent_control_status", "")
        _rc_display = _render_checks(["駆除実施済", "駆除予定", "該当なし"], _rc_val)
        _add_body_paragraph(doc, "【ねずみ・害虫の駆除】", font_size=10, bold=True, space_after=4)
        _add_body_paragraph(doc, f"　{_rc_display}", font_size=10, space_after=6)

    # --- 備考 ---
    _submit_copies = _demo_cfg.get("submit_copies", 2) if _demo_cfg else 2
    _submit_to = _demo_cfg.get("submit_to", "") if _demo_cfg else ""
    _note_parts = [f"提出部数　{_submit_copies}部"]
    if _submit_to:
        _note_parts.append(f"提出先: {_submit_to}")
    _add_body_paragraph(doc, f"備考：{'／'.join(_note_parts)}", font_size=8, space_after=2)

    doc.save(output_path)
    return output_path


# ========== 2. 近隣説明報告書 ==========

def generate_explanation_report(data, output_path):
    """近隣説明報告書を生成
    公式テンプレートがある区はそちらを使用、なければ自作生成
    """
    ward_name = data.get("ward_name", "")

    # 公式テンプレートがあればそちらを使用
    official = _fill_official_report(ward_name, data, output_path)
    if official:
        return official

    # フォールバック: 自作Word生成
    wc = get_ward_config(ward_name)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    _add_heading_paragraph(doc, "近 隣 説 明 報 告 書", font_size=18)
    _add_body_paragraph(doc, f"（{wc['ordinance_name']}{wc['explanation_article']}の規定による報告）",
                        font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    _add_body_paragraph(doc, f"　　　　　　　　　　　　　　　　　　　　　　　　{data.get('submit_date', '令和　年　月　日')}")
    _add_body_paragraph(doc, f"　{ward_name}{wc['suffix']}　殿")
    _add_body_paragraph(doc, "")
    # 報告者情報（右寄せ）
    _add_body_paragraph(doc, f"報告者　住所　{data.get('applicant_address', '')}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_body_paragraph(doc, f"氏名　{data.get('applicant_name', '')}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_body_paragraph(doc, f"電話　{data.get('applicant_tel', '')}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_body_paragraph(doc, "")
    _add_body_paragraph(doc, "　下記建築計画について、近隣関係住民に対し説明を行いましたので報告します。", space_after=12)

    # 建築計画概要
    _add_body_paragraph(doc, "１．建築計画の概要", bold=True, space_after=6)
    overview_rows = [
        ("建築物の名称", data.get("building_name", "")),
        ("建築場所", data.get("site_address", "")),
        ("用途", data.get("building_use", "")),
        ("構造・規模", f"{data.get('structure', '')}　地上{data.get('floors_above', '')}階 地下{data.get('floors_below', '')}階"),
        ("高さ", data.get("height", "") + " m"),
        ("着工予定日", data.get("start_date", "")),
        ("完了予定日", data.get("end_date", "")),
    ]

    table1 = doc.add_table(rows=len(overview_rows), cols=2)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table1)
    for i, (label, value) in enumerate(overview_rows):
        _set_cell(table1.cell(i, 0), label, font_size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(table1.cell(i, 1), value, font_size=10)
        table1.cell(i, 0).width = Cm(4.0)
        table1.cell(i, 1).width = Cm(12.0)

    _add_body_paragraph(doc, "", space_after=6)

    # 説明の実施状況
    _add_body_paragraph(doc, "２．説明の実施状況", bold=True, space_after=6)
    explain_rows = [
        ("説明実施日", data.get("explanation_date", "")),
        ("説明方法", data.get("explanation_method", "個別訪問による説明")),
        ("説明範囲", data.get("explanation_range", f"建築予定地から半径{data.get('radius_m', 50)}m以内の近隣住民")),
        ("説明対象戸数", data.get("target_count", "") + " 戸"),
        ("説明済み戸数", data.get("explained_count", "") + " 戸"),
        ("不在等未説明", data.get("unexplained_count", "") + " 戸"),
    ]

    table2 = doc.add_table(rows=len(explain_rows), cols=2)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table2)
    for i, (label, value) in enumerate(explain_rows):
        _set_cell(table2.cell(i, 0), label, font_size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(table2.cell(i, 1), value, font_size=10)
        table2.cell(i, 0).width = Cm(4.0)
        table2.cell(i, 1).width = Cm(12.0)

    _add_body_paragraph(doc, "", space_after=6)

    # 住民からの意見・要望（長文は改行ごとに段落分割しページ送り対応）
    _add_body_paragraph(doc, "３．近隣関係住民からの意見・要望とその対応", bold=True, space_after=6)
    opinions = data.get("opinions", "特になし")
    opinion_lines = opinions.split("\n") if "\n" in opinions else [opinions]
    for idx, line in enumerate(opinion_lines):
        sa = 12 if idx == len(opinion_lines) - 1 else 2
        _add_body_paragraph(doc, f"　{line}", space_after=sa)

    # 添付書類
    _add_body_paragraph(doc, "添付書類", bold=True, font_size=9)
    _add_body_paragraph(doc, "　１．近隣説明範囲図", font_size=9)
    _add_body_paragraph(doc, "　２．説明配布資料（工事のお知らせ）の写し", font_size=9)

    doc.save(output_path)
    return output_path


# ========== 3. 工事のお知らせ（住民配布チラシ） ==========

def generate_construction_notice(data, output_path):
    """工事のお知らせ（住民配布用チラシ）を生成"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    _add_heading_paragraph(doc, "工 事 の お 知 ら せ", font_size=22)
    _add_body_paragraph(doc, "", space_after=12)

    _add_body_paragraph(doc, "近隣の皆様へ", font_size=12, bold=True, space_after=12)

    # 挨拶文（改行ごとに段落分割しページ送り対応）
    greeting = data.get("greeting_text",
        "平素は格別のご理解を賜り、厚く御礼申し上げます。\n"
        "このたび、下記のとおり工事を実施させていただくこととなりました。\n"
        "工事期間中は、騒音・振動等でご迷惑をおかけいたしますが、\n"
        "安全管理には十分注意して施工いたしますので、何卒ご理解ご協力のほど\n"
        "よろしくお願い申し上げます。"
    )
    greeting_lines = greeting.split("\n") if "\n" in greeting else [greeting]
    for idx, line in enumerate(greeting_lines):
        sa = 16 if idx == len(greeting_lines) - 1 else 2
        _add_body_paragraph(doc, line, font_size=10.5, space_after=sa)

    _add_body_paragraph(doc, "記", font_size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # 工事概要テーブル
    notice_rows = [
        ("工事名称", data.get("site_name", "")),
        ("工事場所", data.get("site_address", "")),
        ("工事内容", data.get("work_content", "")),
        ("工事期間", f"{data.get('start_date', '')} ～ {data.get('end_date', '')}"),
        ("作業時間", data.get("work_hours", "午前8時00分 ～ 午後5時00分")),
        ("休工日", data.get("holidays", "日曜日・祝日")),
        ("施工者", data.get("constructor_name", "")),
        ("現場責任者", data.get("site_manager", "")),
        ("連絡先", data.get("constructor_tel", "")),
    ]

    # 発注者がある場合
    if data.get("client_name"):
        notice_rows.insert(0, ("発注者", data.get("client_name", "")))

    table = doc.add_table(rows=len(notice_rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    for i, (label, value) in enumerate(notice_rows):
        _set_cell(table.cell(i, 0), label, font_size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(table.cell(i, 1), value, font_size=11)
        table.cell(i, 0).width = Cm(3.5)
        table.cell(i, 1).width = Cm(12.5)

    _add_body_paragraph(doc, "", space_after=16)

    # 安全対策
    safety = data.get("safety_measures",
        "・工事車両の出入りには誘導員を配置いたします。\n"
        "・粉塵対策として散水・養生を徹底いたします。\n"
        "・騒音・振動が発生する作業は、事前にお知らせいたします。"
    )
    _add_body_paragraph(doc, "【安全対策について】", font_size=10, bold=True, space_after=4)
    safety_lines = safety.split("\n") if "\n" in safety else [safety]
    for idx, line in enumerate(safety_lines):
        sa = 16 if idx == len(safety_lines) - 1 else 2
        _add_body_paragraph(doc, line, font_size=10, space_after=sa)

    # 問い合わせ先
    _add_body_paragraph(doc, "【お問い合わせ先】", font_size=10, bold=True, space_after=4)
    contact = (
        f"　{data.get('constructor_name', '')}\n"
        f"　現場責任者: {data.get('site_manager', '')}\n"
        f"　電話: {data.get('constructor_tel', '')}"
    )
    _add_body_paragraph(doc, contact, font_size=10, space_after=12)

    # 日付・発行者
    _add_body_paragraph(doc, f"　　　　　　　　　　　　　　　　　　　{data.get('submit_date', '令和　年　月　日')}")
    _add_body_paragraph(doc, f"　　　　　　　　　　　　　　　　　　　{data.get('constructor_name', '')}")

    doc.save(output_path)
    return output_path


# ========== 4. 近隣説明範囲図（Word版） ==========

def generate_map_document(data, map_png_path, output_path, building_pins=None):
    """近隣説明範囲図をWord文書として生成（地図画像埋め込み + 建物リスト）"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(29.7)  # A4横
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    _add_heading_paragraph(doc, "近 隣 説 明 範 囲 図", font_size=18)
    _add_body_paragraph(doc, "", space_after=4)

    # 地図画像
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(map_png_path, width=Cm(24.0))

    _add_body_paragraph(doc, "", space_after=4)

    # 情報テーブル
    info_rows = [
        ("工事名称", data.get("site_name", "")),
        ("工事場所", data.get("site_address", "")),
        ("説明範囲", f"工事現場から半径 {data.get('radius_m', 50)}m 以内"),
        ("施工者", data.get("constructor_name", "")),
    ]
    table = doc.add_table(rows=len(info_rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)
    for i, (label, value) in enumerate(info_rows):
        _set_cell(table.cell(i, 0), label, font_size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(table.cell(i, 1), value, font_size=10)
        table.cell(i, 0).width = Cm(3.5)
        table.cell(i, 1).width = Cm(22.0)

    # 建物番号リスト（ピンがある場合は必ず出力、ラベル空欄でもOK）
    if building_pins:
        _add_body_paragraph(doc, "", space_after=4)
        _add_body_paragraph(doc, "近隣建物一覧", font_size=12, bold=True, space_after=6)

        bld_table = doc.add_table(rows=len(building_pins) + 1, cols=3)
        bld_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(bld_table)

        # ヘッダー行
        _set_cell(bld_table.cell(0, 0), "番号", font_size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(bld_table.cell(0, 1), "建物名称・用途", font_size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(bld_table.cell(0, 2), "備考", font_size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        bld_table.cell(0, 0).width = Cm(2.0)
        bld_table.cell(0, 1).width = Cm(12.0)
        bld_table.cell(0, 2).width = Cm(11.5)

        for i, pin in enumerate(building_pins):
            row_idx = i + 1
            _set_cell(bld_table.cell(row_idx, 0), str(pin["no"]), font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell(bld_table.cell(row_idx, 1), pin.get("label", ""), font_size=10)
            _set_cell(bld_table.cell(row_idx, 2), "", font_size=10)
            bld_table.cell(row_idx, 0).width = Cm(2.0)
            bld_table.cell(row_idx, 1).width = Cm(12.0)
            bld_table.cell(row_idx, 2).width = Cm(11.5)

    doc.save(output_path)
    return output_path
